import requests
from bs4 import BeautifulSoup
import time
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import wave
import os

user_agent = {"User-agent" : "Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:69.0) Gecko/20100101 Firefox/69.0"}

referat_file = open("essay.txt","w",encoding="utf-8") 
referat_document =  Document()

#data is the entire text from html page 
def create_essay(data,country):
    heading = referat_document.add_heading(country, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    index = 0
    paragraphs = data.find_all('p')
    for par in paragraphs:
        if (index % 3) == 0:
            #print(par.get_text())
            text = par.get_text() 
            text = str(text)
            if(len(text)):
                referat_file.writelines(text)
                referat_document.add_paragraph(text)
        index = index + 1
    referat_file.close()
    referat_document.save(country+".docx")
    os.startfile(country+".docx")

def get_Country_Information():
    name = input("Introdu numele tau pentru a face cunostiinta cu ReferatBot ")
    print("ReferatBot te saluta , draga " + name)
    country = input("Introdu numele tarii despre care trebuie sa faci referat : ")
    URL = "https://ro.wikipedia.org/wiki/" + country
    page = requests.get(URL,headers=user_agent)
    data = BeautifulSoup(page.content,"html.parser")

    create_essay(data , country)
    #test_create_Document()
    print("The essay is getting ready")
    #time.sleep(5)

    print("The essay is done . Good luck !! ;)")
    #time.sleep(5)

get_Country_Information()    
